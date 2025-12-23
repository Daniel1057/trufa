from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.select import Select
import pandas as pd
import docx
from urllib.request import urlretrieve
import os
from docx.enum.section import WD_ORIENTATION
import cv2
import time


def image_resize(image, width = None, height = None, inter = cv2.INTER_AREA):
    # initialize the dimensions of the image to be resized and
    # grab the image size
    dim = None
    img = cv2.imread('./brp/img/'+image)
    img =cv2.rotate(img,cv2.ROTATE_90_CLOCKWISE)
    h,w,c = img.shape

    if(w<width/2):
        width = h

    # check to see if the width is None
    if width is None:
        # calculate the ratio of the height and construct the
        # dimensions
        r = height / float(h)
        dim = (int(w * r), height)

    # otherwise, the height is None
    else:
        # calculate the ratio of the width and construct the
        # dimensions
        r = width / float(w)
        dim = (width, int(h * r))

    # resize the image
    resized = cv2.resize(img, dim, interpolation = inter)

    # return the resized image
    cv2.imwrite('./brp/img/'+image,resized)
    #return resized
def descarga_img(enlace,titulo):
    dir = os.path.abspath('./brp/img/')
    fichero = titulo+".png"
    work_path=os.path.join(dir,fichero)
    urlretrieve(enlace,work_path)
    return(fichero)

def genera_word(titulo):
    document = docx.Document()
    section = document.sections[-1]
    section.orientation = WD_ORIENTATION.PORTRAIT
    document.add_picture('endurorecambios.png')
    footer = section.footer
    footer.add_paragraph(titulo)

    document.save('./brp/documento/'+titulo +'.docx')

def add_word(titulo,fichero,imagen,df):
    document = docx.Document('./brp/documento/'+fichero +'.docx')
    document.add_page_break()
    #section = document.sections[0]
    #header = section.header
    #header.para = header.paragraphs[0]
    #header_para.text = titulo
    encabezado = document.add_heading (titulo,2)
    encabezado.alignment = 1
    image_resize(imagen, width=400)
    document.add_picture('./brp/img/'+imagen)
    document.add_page_break()
    encabezado = document.add_heading(titulo, 2)
    encabezado.alignment = 1
    t=document.add_table(len(df[0]['#'].tolist())+1,4)
    t.rows[0].cells[0].text = '#'
    t.rows[0].cells[1].text = 'Part #'
    t.rows[0].cells[2].text = 'Description'
    t.rows[0].cells[3].text = 'Qty'
    columna1 = []
    columna2 = []
    columna3 = []
    columna4 = []
    zcol1 = 1
    zcol2 = 1
    zcol3 = 1
    zcol4 = 1
    for i in range(len(df)):
        columna1 = df[i]['#'].tolist()
        columna2 = df[i]['Part #'].tolist()
        columna3 = df[i]['Description'].tolist()
        columna4 = df[i]['Qty'].tolist()
        for valor1 in columna1:
            cell = t.cell(zcol1, 0)
            cell.text = str(valor1)
            zcol1 = zcol1+1
        for valor2 in columna2:
            cell = t.cell(zcol2, 1)
            cell.text = str(valor2)
            zcol2 = zcol2+1
        for valor3 in columna3:
            cell = t.cell(zcol3, 2)
            cell.text = str(valor3)
            zcol3 = zcol3+1
        for valor4 in columna4:
            cell = t.cell(zcol4, 3)
            cell.text = str(valor4)
            zcol4 = zcol4+1
    t.style = 'Colorful List'
    t.autofit = False
    sections = document.sections
    for section in sections:
        section.orientation = WD_ORIENTATION.LANDSCAPE
    document.save('./brp/documento/'+fichero+'.docx')

def genera_fichero_a(fichero):
    cadena = fichero
    cadena=cadena[cadena.find ('assembly'):len(cadena)]
    cadena = cadena[cadena.find('<a'):len(cadena)]
    enlaces = '\n'
    while (cadena.find ('<a')>=0):
        cadena = cadena[cadena.find('>')+1:len(cadena)]
        subcadena = cadena[:cadena.find('</')]
        cadena = cadena[cadena.find('<a'):]
        enlaces = enlaces + subcadena + '\n'
        print (enlaces)
    enlaces = enlaces.replace('  ',' ')
    return enlaces



def inicio(anyo,tipo,modelo,product):
    driver = webdriver.Firefox()
    driver.get('http://epc.brp.com/')

    select_element = driver.find_element(By.ID, 'ctl00_ddlCatalog')
    select_object = Select(select_element)
    select_object.select_by_visible_text(str(product))
    driver.find_elements_by_tag_name('iframe')
    driver.switch_to.frame('ctl00_MainPlaceHolder_AssemblyTreeIFrm')
    #driver.find_element_by_link_text('2019').click()
    driver.find_element_by_link_text(str(anyo)).click()
    #driver.find_element_by_link_text('140 - SSV - International - Maverick Series').click()
    driver.find_element_by_link_text(str(tipo)).click()
    #modelo = '003 - Maverick Turbo R - International, 2019'
    driver.find_element_by_link_text(modelo).click()
    time.sleep(1)
    enlaces = genera_fichero_a(driver.find_element_by_xpath("//*[@id='treeContainerDiv']").get_attribute("outerHTML"))
    fichero = modelo
    #genera_word(fichero)
    entrar = True
    for line in enlaces.splitlines():
        if(line=='10- Electrical - Engine Harness And Electronic Module'):
            entrar = True
        if (entrar and line != ''):
            try:
                print(line)
                familia = line
                time.sleep(1)
                elemento = driver.find_element_by_partial_link_text(familia)
                elemento.click()
                titulo = familia
                #driver.implicitly_wait(300)
                time.sleep(1)
                driver.switch_to.parent_frame()
                driver.switch_to.frame('ctl00_MainPlaceHolder_AssemblyDetailImageIFrm')
                time.sleep(2)
                driver.find_element_by_xpath('//*[@id="imgZoomIn"]').click()
                driver.find_element_by_xpath('//*[@id="imgZoomIn"]').click()
                driver.find_element_by_xpath('//*[@id="imgZoomIn"]').click()
                elemento=driver.find_element_by_tag_name("img")
                imagen = elemento.get_attribute("src")
                print(imagen)
                fichero_img=descarga_img(imagen,titulo)
                driver.implicitly_wait(100)
                driver.switch_to.parent_frame()
                driver.switch_to.frame('AssemblyDetailPartListIFrm')
                tabla = pd.read_html(driver.page_source)
                error = tabla
                driver.switch_to.parent_frame()
                driver.switch_to.frame('ctl00_MainPlaceHolder_AssemblyTreeIFrm')
                add_word(titulo,fichero,fichero_img,tabla)
            except:
                error= ('No hay tabla para:' + line)
                raise
            print(error)


if __name__ == '__main__':

    inicio()

