import pandas as pd
import pandas_datareader.data as web
import base64
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
import time
import os
from docx.enum.section import WD_ORIENTATION
import docx
import cv2
from sqlalchemy import create_engine
import pymysql
import prestaweb
import Alim_bbdd_sah
from os import listdir

def image_resize(image, width = None, height = None, inter = cv2.INTER_AREA):
    # initialize the dimensions of the image to be resized and
    # grab the image size
    dim = None
    img = cv2.imread('./beta/img/'+image)
    img =cv2.rotate(img,cv2.ROTATE_90_CLOCKWISE)
    h,w,c = img.shape

    if(w<width/2):
        width = h

    # check to see if the width is None
    if width is None:
        # calculate the ratio of the height and construct the
        # dimensions
        r = height / float(h)
        dim = (int(w * r), height)

    # otherwise, the height is None
    else:
        # calculate the ratio of the width and construct the
        # dimensions
        r = width / float(w)
        dim = (width, int(h * r))

    # resize the image
    resized = cv2.resize(img, dim, interpolation = inter)

    # return the resized image
    cv2.imwrite('./beta/img/'+image,resized)
    #return resized

def genera_word(titulo):
    document = docx.Document()
    section = document.sections[-1]
    section.orientation = WD_ORIENTATION.PORTRAIT
    document.add_picture('endurorecambios.png')
    footer = section.footer
    titulo = titulo.replace(' ','_')
    footer.add_paragraph(titulo)

    document.save('./beta/documento/'+titulo +'.docx')
def busca_descripcion(valor,tarifa):

    print (tarifa)
    if (valor in tarifa.values):
        tar = tarifa['Article'] == valor
        print(tar)
        fila = tarifa[tar]['Descripcion']
        return(fila)
    else:
        valor = valor.replace('.', '')
        if (valor in tarifa.values):
            tar = tarifa['cod'] == valor
            print(tar)
            fila = tarifa[tar]['Descripcion']
            return (fila)
        else:
            return('')

def add_word(titulo,fichero,imagen,df,tarifa):
    document = docx.Document('./beta/documento/'+titulo +'.docx')
    document.add_page_break()
    #section = document.sections[0]
    #header = section.header
    #header.para = header.paragraphs[0]
    #header_para.text = titulo
    encabezado = document.add_heading (titulo,2)
    encabezado.alignment = 1
    image_resize(imagen, width=400)
    document.add_picture('./beta/img/'+imagen)
    document.add_page_break()
    encabezado = document.add_heading(titulo, 2)
    encabezado.alignment = 1
    t=document.add_table(len(df[1]['Ref.'].tolist())+1,4)
    t.rows[0].cells[0].text = 'Ref.'
    t.rows[0].cells[1].text = 'Part Number'
    t.rows[0].cells[2].text = 'Description'
    t.rows[0].cells[3].text = 'Part'
    columna1 = []
    columna2 = []
    columna3 = []
    columna4 = []
    zcol1 = 1
    zcol2 = 1
    zcol3 = 1
    zcol4 = 1
    i=1
    columna1 = df[i]['Ref.'].tolist()
    columna2 = df[i]['Part Number'].tolist()
    #columna3 = df[i]['Description'].tolist()
    columna4 = df[i]['Parts'].tolist()
    for valor1 in columna1:
        cell = t.cell(zcol1, 0)
        cell.text = str(valor1)
        zcol1 = zcol1+1
    for valor2 in columna2:
        cell = t.cell(zcol2, 1)
        cell.text = str(valor2)
        zcol2 = zcol2+1
    for valor3 in columna2:
        cell = t.cell(zcol3, 2)
        valor3 = busca_descripcion(valor3,tarifa)
        for v_valor3 in valor3:
            cell.text = str(v_valor3)
        zcol3 = zcol3+1
    for valor4 in columna4:
        cell = t.cell(zcol4, 3)
        cell.text = str(valor4)
        zcol4 = zcol4+1
    t.style = 'Colorful List'
    t.autofit = False
    sections = document.sections
    for section in sections:
        section.orientation = WD_ORIENTATION.LANDSCAPE
    document.save('./beta/documento/'+titulo+'.docx')

def carga_selector_moto():
    moto = ['ENDURO RR 125 2T MY22','ENDURO RR 200 2T MY22','ENDURO RR 250 2T MY22','ENDURO RR 300 2T MY22','ENDURO RR 350 4T MY22','ENDURO RR 390 4T MY22','ENDURO RR 430 4T MY22','ENDURO RR 480 4T MY22','ENDURO XTRAINER 250 2T MY22','ENDURO XTRAINER 300 2T MY22']
    #moto = ['ENDURO RR 350 4T MY22','ENDURO RR 390 4T MY22','ENDURO RR 430 4T MY22', 'ENDURO RR 480 4T MY22','ENDURO XTRAINER 250 2T MY22','ENDURO XTRAINER 300 2T MY22','ENDURO XTRAINER 300 2T MY22']
    return moto

def despiece_beta(driver,moto,componente,tarifa):
    titulo = moto
    time.sleep(8)
    modulo = driver.find_element_by_css_selector('html body#body div#app div#wrap div#portal-columns div.row div#content_page.col-xl-12 div.noStampa div.drawToolbar.row div.col-lg-10.col-md-9.col-sm-8.row div.col-lg-4.col-sm-6.mb-2.pl-0.drawToolbarTitles').text
    grupo = driver.find_element_by_css_selector('div.col-lg-4:nth-child(2)').text
    grupo = grupo.replace('Group:\n', '')
    #moto = moto.replace(' ','_')
    fichero =moto + '_'+ grupo+'_'+str(componente)
    tabla = pd.read_html(driver.page_source)
    #tabla = web.DataReader(driver.page_source)
    print(tabla[1])
    #driver.find_element_by_css_selector("#immagine_tavola").click()
    #driver.find_element_by_css_selector("#immagine_tavola").click()
    #driver.find_element_by_css_selector("#immagine_tavola").click()
    canvas = driver.find_element_by_css_selector("#immagine_tavola")
    # get the canvas as a PNG base64 string
    canvas_base64 = driver.execute_script("return arguments[0].toDataURL('image/png').substring(21);", canvas)
    # decode
    canvas_png = base64.b64decode(canvas_base64)
    print (os.getcwd())
    # save to a file
    with open('./beta/img/'+fichero+'.png', 'wb') as f:
        f.write(canvas_png)

    driver.find_element_by_xpath('/html/body/div[1]/div/div[3]/div[2]/div/div/div[1]/div[1]/a').send_keys(Keys.ENTER)
    add_word(moto, titulo+'.docx', fichero+'.png', tabla,tarifa)
    time.sleep(1)
    tabla[1]['Categoria'] = grupo
    return tabla[1]

def cargar_df(moto,df_dataframe):
    sqlEngine = create_engine('mysql+pymysql://Minemas_R:b1EppZu9bX6vhvMf@endurorecambios.com/Alim')
    dbConnection = sqlEngine.connect()
    df_dataframe.to_sql(name=moto, con=dbConnection, if_exists='replace', index=False)
    dbConnection.close()

def inicio_beta(modelos):
    #motos = list(modelos.split(","))
    #motos = carga_selector_moto()
    motos = modelos
    tarifa = pd.read_excel('TARIFA 15-09-2021.xlsx')
    tarifa['cod'] = tarifa['Article'].str.replace('.', '', regex=True)
    driver = webdriver.Firefox()
    driver.get('https://partsfinder.softway.it/beta/')
    time.sleep(3)
   # if depurando:
   #     motos = carga_selector_moto()

    v_df_to_sql = False
    v_df_for_moto = False
    df_moto = pd.DataFrame()
    for moto in motos:
        #if(v_df_to_sql):
            #cargar_df (v_moto,df_moto)
        v_moto = moto.replace(' ','_')
        print ('--------------------'+v_moto+'--------------------------')
        genera_word(moto)

        driver.find_element_by_css_selector('.select2').click()
        driver.find_element_by_css_selector('.select2-search__field').send_keys(moto)
        driver.find_element_by_css_selector('.select2-search__field').send_keys(Keys.ENTER)
        time.sleep(8)
        driver.find_element_by_xpath('/html/body/div[1]/div/div[3]/div[2]/div/div/div/div[3]/div/div[1]/span/span[1]/span').click()
        time.sleep(8)
        repeticiones = driver.find_element_by_xpath('/html/body/span/span/span[2]/ul/li[1]/ul').get_attribute('outerHTML').count('</li>')
        repeticiones_2 = driver.find_element_by_xpath('/html/body/span/span/span[2]/ul/li[2]/ul').get_attribute('outerHTML').count('</li>')
        repeticiones = repeticiones_2 + repeticiones
        print(repeticiones)
        entrar = True
        for componente in range (repeticiones):
            print(componente)
            if(entrar):
                if(componente>0):
                    driver.find_element_by_xpath('/html/body/div[1]/div/div[3]/div[2]/div/div/div/div[3]/div/div[1]/span/span[1]/span').click()
                    driver.find_element_by_xpath('/html/body/div[1]/div/div[3]/div[2]/div/div/div/div[3]/div/div[1]/span/span[1]/span').send_keys(Keys.ARROW_DOWN)

                driver.find_element_by_xpath('/html/body/div[1]/div/div[3]/div[2]/div/div/div/div[3]/div/div[1]/span/span[1]/span').send_keys(Keys.ENTER)

                time.sleep(4)
                driver.find_element_by_xpath('/html/body/div[1]/div/div[3]/div[2]/div/div/div/div[3]/div/a').click()
                time.sleep(4)
                # Solo para depurardriver.find_element_by_xpath('/html/body/div[1]/div/div[3]/div[2]/div/div/div[1]/div[1]/a').send_keys(Keys.ENTER)

                v_temp_df_moto = despiece_beta(driver, v_moto, componente,tarifa)
                v_temp_df_moto['modelo'] = moto
                v_temp_df_moto['componente'] = v_moto + '_' + str(componente)
                if(v_df_to_sql==False):
                    v_df_to_sql = True
                    if(v_df_for_moto==False):
                        v_df_for_moto = True
                        df_moto = v_temp_df_moto
                else:
                    df_moto = pd.concat([df_moto,v_temp_df_moto],axis=0)
    df_moto.to_excel('beta_todas_referencias.xlsx')
    Alim_bbdd_sah.carga_tarifa_bbdd(df_moto,'comparativa_beta')
    df_moto = df_moto.drop_duplicates

def carga_df(df_moto):
    tarifa = pd.read_excel('TARIFA 15-09-2021.xlsx')
    tarifa['cod'] = tarifa['Article'].str.replace('.', '', regex=True)
    if('NEW' in df_moto.values):
        df_moto = df_moto[df_moto['Description'] == 'NEW']
        df_completo = pd.merge(left=df_moto, right=tarifa, how='left', left_on='Part Number', right_on='Article')
        print(df_completo)
        df_completo = df_completo.drop(['Unnamed: 0'], axis=1)
        df_completo = df_completo.drop(['Unnamed: 5'], axis=1)
        df_completo['COD_PROD'] = df_completo['Part Number'].str.replace('.', '', regex=True)
        df_completo = pd.merge(left=df_completo, right=tarifa, how='left', left_on='COD_PROD', right_on='cod')
        df_completo["Descripcion_x"].fillna(df_completo["Descripcion_y"], inplace=True)
        df_completo["Preu_x"].fillna(df_completo["Preu_y"], inplace=True)
        df_completo["Descuento_x"].fillna(df_completo["Descuento_y"], inplace=True)

        # referencias obsoletas
        df_completo.to_excel('df_completo_new.xlsx')
        if (df_completo[df_completo['Descripcion_x'].str.contains('#', na=False, regex=True)].size > 0):
            df_obsoletos = df_completo[df_completo['Descripcion_x'].str.contains('#', na=False, regex=True)]
            df_completo=df_completo.drop(df_completo[df_completo['Descripcion_x'].str.contains('#', na=False, regex=True)].index)
            df_obsoletos.to_excel('beta_obsoleto.xlsx')
            cargar_obsoletos(df_obsoletos, tarifa, 0)
        prestaweb.ps_web_catego_producto(df_completo,'13')
        return(1)

def cargar_prod_xlsx(fichero):
    df_completo=pd.read_excel (fichero)
    if (df_completo[df_completo['Descripcion_x'].str.contains('#', na=False, regex=True)].size > 0):
        df_obsoletos = df_completo[df_completo['Descripcion_x'].str.contains('#', na=False, regex=True)]
        df_completo = df_completo.drop(
            df_completo[df_completo['Descripcion_x'].str.contains('#', na=False, regex=True)].index)
        df_completo.to_excel('obsoleto.xlsx')
    prestaweb.ps_web_catego_producto(df_completo, '13')
    return(1)

def cargar_obsoletos_xlsx(fichero,fichero_tarifa,iteracion):
    df_completo = pd.read_excel(fichero)
    tarifa = pd.read_excel(fichero_tarifa)
    tarifa['cod'] = tarifa['Article'].str.replace('.', '', regex=True)
    cargar_obsoletos(df_completo, tarifa,iteracion)

def cargar_obsoletos(df_completo, tarifa, iteracion):
    if (df_completo[df_completo['Descripcion_x'].str.contains('#', na=False, regex=True)].size > 0):
        df_obsoletos = df_completo[df_completo['Descripcion_x'].str.contains('#', na=False, regex=True)]
        df_obsoletos = df_obsoletos[['Ref.', 'Part Number', 'Description', 'Parts', 'Categoria', 'modelo','componente','Descripcion_x']]
        df_obsoletos.columns = ['Ref.', 'Part Number', 'Description', 'Parts', 'Categoria', 'modelo', 'componente','new_ref']
        df_obsoletos['old_ref'] = df_obsoletos['Part Number']
        df_obsoletos['new_ref'] = df_obsoletos['new_ref'].str.replace('#', '', regex=True)
        df_obsoletos['Part Number'] = df_obsoletos['new_ref']
        df_obsoletos['COD_PROD'] = df_obsoletos['new_ref'].str.replace('.', '', regex=True)
        df_obsoletos = pd.merge(left=df_obsoletos, right=tarifa, how='left', left_on='Part Number', right_on='Article')
        df_obsoletos = pd.merge(left=df_obsoletos, right=tarifa, how='left', left_on='COD_PROD', right_on='cod')
        df_obsoletos["Descripcion_x"].fillna(df_obsoletos["Descripcion_y"], inplace=True)
        df_obsoletos["Preu_x"].fillna(df_obsoletos["Preu_y"], inplace=True)
        df_obsoletos["Descuento_x"].fillna(df_obsoletos["Descuento_y"], inplace=True)
        df_obsoletos['Descripcion_x'] = df_obsoletos['Descripcion_x'] + ' - ' + df_obsoletos['old_ref']
        df_obsoletos['Descripcion_x'] = df_obsoletos['Descripcion_x'].str.replace('.', '', regex=True)
        df_completo = df_obsoletos
        if (df_completo[df_completo['Descripcion_x'].str.contains('#', na=False, regex=True)].size > 0):
            df_obsoletos = df_completo[df_completo['Descripcion_x'].str.contains('#', na=False, regex=True)]
            df_completo=df_completo.drop(df_completo[df_completo['Descripcion_x'].str.contains('#', na=False, regex=True)].index)
            if(iteracion<4):
                iteracion = iteracion + 1
                cargar_obsoletos(df_completo, tarifa, iteracion)
        prestaweb.ps_web_catego_producto(df_completo,'13')
        print(df_obsoletos)
def recorrer_directorio (ruta):
    arch = listdir(str(ruta))
    ruta_completa = os.getcwd()
    ruta_completa = ruta_completa.replace('/prestaweb.py','')
    ruta_completa = ruta_completa+ruta.replace('.','')
    print(ruta_completa)
    for archivo in arch:
        if (archivo.find('xlsx')>0):
            print(ruta_completa+archivo)
            carga_df(pd.read_excel(ruta_completa+archivo))


if __name__ == '__main__':
    #inicio_beta('ENDURO RR 125 2T MY22')
    #cargar_prod_xlsx('df_completo_new_1.xlsx')
    #cargar_obsoletos_xlsx('df_completo_new_1.xlsx','TARIFA 15-09-2021.xlsx',0)
    recorrer_directorio('./despieces/')