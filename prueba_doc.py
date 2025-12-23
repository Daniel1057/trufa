import docx
from docx.shared import Cm
from docx.enum.section import WD_ORIENTATION



document = docx.Document()



#genera cabecera
document.add_heading ('Titulo principal',0)

#parrafo
p=document.add_paragraph('Esto es un parrafo')
p.add_run('añadimos al parrafo ')
p.add_run('negrita').bold = True

#añadir imagen
document.add_heading('Imágenes',level=1)
document.add_picture('prueba.png',width=Cm(5))

#Tablas
document.add_heading('Tablas',level = 1)
data= (('Manzana',12),('Pera',5),('Naranja',12))

table = document.add_table(rows=1,cols=2)
table.rows[0].cells[0].text = 'Fruta'
table.rows[0].cells[1].text = 'Cantidad'

#Insertar datos
for prod,numbr in data:
    row_cells = table.add_row().cells
    row_cells[0].text = prod
    row_cells[1].text = str(numbr)

document.save('ejemplo.docx')

